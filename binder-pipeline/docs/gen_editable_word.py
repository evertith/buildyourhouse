#!/usr/bin/env python3
"""Editable Word (.docx) companions for Section 2 of the Owner-Builder Job Site Binder.

Reproduces the contract language from the printed PDFs as genuinely editable Word
documents: real tables, real borders, no runs of underscore characters.

Conventions used throughout:
  * Discrete labelled data (Project Address, Check Number, ...) -> 2-column field
    tables whose value cell carries a bottom border.
  * Blanks that sit inside a sentence of contract prose -> bold bracketed
    placeholders, e.g. [AMOUNT], which stay readable and are trivially typed over.
  * Free-text areas -> ruled writing lines built from bordered table rows.

Output: out/editable-documents/word/
"""

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

# ---------------------------------------------------------------------------
# Brand constants (mirrors design.py used by the PDF pipeline)
# ---------------------------------------------------------------------------

HERE = os.path.dirname(os.path.abspath(__file__))
PIPELINE = os.path.dirname(HERE)
OUT_DIR = os.path.join(PIPELINE, "out", "editable-documents", "word")

BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)

INK = RGBColor(0x00, 0x00, 0x00)
FURNITURE_GREY = RGBColor(0x59, 0x59, 0x59)
LABEL_GREY = RGBColor(0x40, 0x40, 0x40)

HEADER_FILL = "E6E6E6"   # table header band
SUBTOTAL_FILL = "F2F2F2"  # subtotal / section band
NOTE_FILL = "F7F2E7"      # warning box (amber)
NOTICE_FILL = "F2F2F2"    # notice box (neutral — informational, not a caution)

COPYRIGHT = "© 2026 Build Your House · build-your-house.com"
DISCLAIMER = (
    "Template for general reference — have your attorney review before use. "
    "Not legal advice."
)

CONTENT_WIDTH = 7.0  # inches, letter page less 0.75" margins each side
CHECKBOX = "☐"


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------

_EDGE_ORDER = ("top", "left", "bottom", "right", "insideH", "insideV")


def _border_el(tag, val, sz, color):
    el = OxmlElement("w:" + tag)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    return el


def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    """Apply per-edge borders. Each edge is None (leave), False (nil) or a
    (size_eighths, hex_color) tuple."""
    spec = {"top": top, "left": left, "bottom": bottom, "right": right}
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in _EDGE_ORDER:
        value = spec.get(edge)
        if value is None:
            continue
        if value is False:
            borders.append(_border_el(edge, "nil", 0, "auto"))
        else:
            sz, color = value
            borders.append(_border_el(edge, "single", sz, color))
    if len(borders):
        tcPr.append(borders)


def clear_table_borders(table):
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in _EDGE_ORDER:
        borders.append(_border_el(edge, "nil", 0, "auto"))
    tblPr.append(borders)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def fixed_layout(table, widths):
    """Pin the table to a fixed layout with explicit column widths (inches)."""
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = Inches(widths[idx])
    for idx, col in enumerate(table.columns):
        if idx < len(widths):
            col.width = Inches(widths[idx])


def set_row_height(row, points, exact=False):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(points * 20)))
    trHeight.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(trHeight)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def set_run_font(run, name, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------


class BinderDoc:
    """Thin wrapper over python-docx that enforces the binder's house style."""

    def __init__(self, title):
        self.doc = Document()
        self.title = title
        self._configure_styles()
        self._configure_page()
        self._configure_footer()
        self.doc.core_properties.title = title
        self.doc.core_properties.author = "Build Your House"
        self.doc.core_properties.comments = COPYRIGHT

    # -- setup -------------------------------------------------------------

    def _configure_styles(self):
        styles = self.doc.styles

        normal = styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = BODY_SIZE
        normal.font.color.rgb = INK
        rPr = normal.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), BODY_FONT)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.line_spacing = 1.12

        heading_specs = {
            "Heading 1": (18, Pt(0), Pt(10)),
            "Heading 2": (13, Pt(14), Pt(6)),
            "Heading 3": (11.5, Pt(10), Pt(4)),
        }
        for name, (size, before, after) in heading_specs.items():
            style = styles[name]
            style.font.name = BODY_FONT
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = INK
            style.paragraph_format.space_before = before
            style.paragraph_format.space_after = after
            style.paragraph_format.keep_with_next = True
            hrPr = style.element.get_or_add_rPr()
            hrFonts = hrPr.find(qn("w:rFonts"))
            if hrFonts is None:
                hrFonts = OxmlElement("w:rFonts")
                hrPr.insert(0, hrFonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                hrFonts.set(qn(attr), BODY_FONT)

        for name in ("List Bullet", "List Number"):
            style = styles[name]
            style.font.name = BODY_FONT
            style.font.size = BODY_SIZE
            style.paragraph_format.space_after = Pt(3)

    def _configure_page(self):
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.7)
        section.footer_distance = Inches(0.4)

    def _configure_footer(self):
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            para = footer.paragraphs[0]
            para.text = ""
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(COPYRIGHT)
            set_run_font(run, BODY_FONT, Pt(8), color=FURNITURE_GREY)

    # -- text --------------------------------------------------------------

    def _emit_runs(self, para, text, bold=False, italic=False, size=None,
                   color=None):
        """Write text into a paragraph, bolding [BRACKETED] fill-in placeholders."""
        chunk = ""
        idx = 0
        while idx < len(text):
            ch = text[idx]
            if ch == "[":
                close = text.find("]", idx)
                if close != -1:
                    if chunk:
                        run = para.add_run(chunk)
                        set_run_font(run, BODY_FONT, size, bold, italic, color)
                        chunk = ""
                    run = para.add_run(text[idx:close + 1])
                    set_run_font(run, BODY_FONT, size, True, italic, color)
                    idx = close + 1
                    continue
            chunk += ch
            idx += 1
        if chunk:
            run = para.add_run(chunk)
            set_run_font(run, BODY_FONT, size, bold, italic, color)

    def h1(self, text):
        para = self.doc.add_paragraph(style="Heading 1")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._emit_runs(para, text, bold=True, size=Pt(18))
        return para

    def subtitle(self, text):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(14)
        self._emit_runs(para, text, bold=True, size=Pt(12), color=FURNITURE_GREY)
        return para

    def h2(self, text):
        para = self.doc.add_paragraph(style="Heading 2")
        self._emit_runs(para, text, bold=True, size=Pt(13))
        self._rule_under(para)
        return para

    def h3(self, text):
        para = self.doc.add_paragraph(style="Heading 3")
        self._emit_runs(para, text, bold=True, size=Pt(11.5))
        return para

    def _rule_under(self, para):
        pPr = para._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        borders.append(_border_el("bottom", "single", 6, "8C8C8C"))
        pPr.append(borders)

    def p(self, text, bold=False, italic=False, size=None, color=None,
          space_after=Pt(6)):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = space_after
        self._emit_runs(para, text, bold=bold, italic=italic, size=size,
                        color=color)
        return para

    def label(self, text):
        return self.p(text, bold=True, space_after=Pt(3))

    def bullets(self, items):
        for item in items:
            para = self.doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_after = Pt(2)
            self._emit_runs(para, item)

    def numbered(self, items):
        for item in items:
            para = self.doc.add_paragraph(style="List Number")
            para.paragraph_format.space_after = Pt(4)
            self._emit_runs(para, item)

    def checklist(self, items, indent=0.0):
        for item in items:
            para = self.doc.add_paragraph()
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.left_indent = Inches(indent + 0.3)
            para.paragraph_format.first_line_indent = Inches(-0.3)
            box = para.add_run(CHECKBOX + "  ")
            set_run_font(box, "Segoe UI Symbol", Pt(12))
            self._emit_runs(para, item)

    def page_break(self):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.add_run().add_break(WD_BREAK.PAGE)

    def spacer(self, points=6):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(points)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run("")
        set_run_font(run, BODY_FONT, Pt(2))

    # -- form furniture ----------------------------------------------------

    def fields(self, rows, label_width=1.9):
        """Labelled fill-in fields.

        `rows` is a list of rows; each row is a list of one or two labels.
        A single label spans the full content width; two labels split it.
        """
        value_width = (CONTENT_WIDTH - 2 * label_width) / 2
        widths = [label_width, value_width, label_width, value_width]

        table = self.doc.add_table(rows=len(rows), cols=4)
        clear_table_borders(table)
        fixed_layout(table, widths)

        for r_idx, labels in enumerate(rows):
            row = table.rows[r_idx]
            set_row_height(row, 26)
            cells = row.cells
            for cell in cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)

            if len(labels) == 1:
                value_cell = cells[1].merge(cells[3])
                pairs = [(cells[0], value_cell, labels[0])]
            else:
                pairs = [(cells[0], cells[1], labels[0]),
                         (cells[2], cells[3], labels[1])]

            for label_cell, value_cell, text in pairs:
                run = label_cell.paragraphs[0].add_run(text + ":")
                set_run_font(run, BODY_FONT, Pt(10.5), bold=True,
                             color=LABEL_GREY)
                set_cell_borders(value_cell, bottom=(6, "8C8C8C"))
        self.spacer(4)
        return table

    def ruled(self, count, width=CONTENT_WIDTH, pitch=26):
        """Blank ruled writing lines (bordered empty table rows)."""
        table = self.doc.add_table(rows=count, cols=1)
        clear_table_borders(table)
        fixed_layout(table, [width])
        for row in table.rows:
            set_row_height(row, pitch)
            cell = row.cells[0]
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            set_cell_borders(cell, bottom=(6, "B3B3B3"))
        self.spacer(6)
        return table

    def grid(self, headers, rows, widths, bold_rows=(), band_rows=(),
             align_right=()):
        """A real Word table with a shaded, repeating header row."""
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = self.doc.styles["Table Grid"]
        fixed_layout(table, widths)

        head = table.rows[0]
        repeat_header(head)
        set_row_height(head, 22)
        for idx, text in enumerate(headers):
            cell = head.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            if idx in align_right:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(text)
            set_run_font(run, BODY_FONT, Pt(10), bold=True)
            shade_cell(cell, HEADER_FILL)

        for r_idx, values in enumerate(rows):
            row = table.rows[r_idx + 1]
            set_row_height(row, 26)
            is_bold = r_idx in bold_rows
            is_band = r_idx in band_rows or is_bold
            for c_idx, value in enumerate(values):
                cell = row.cells[c_idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = Pt(0)
                if c_idx in align_right:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if value:
                    run = para.add_run(value)
                    set_run_font(run, BODY_FONT, Pt(10.5), bold=is_bold)
                if is_band:
                    shade_cell(cell, SUBTOTAL_FILL)
        self.spacer(6)
        return table

    def signature_block(self, columns, gutter=0.4):
        """Side-by-side signature panels.

        `columns` is a list of (heading, [field labels]) tuples.
        """
        n = len(columns)
        depth = max(len(fields) for _, fields in columns)
        col_width = (CONTENT_WIDTH - gutter * (n - 1)) / n
        widths = []
        for i in range(n):
            widths.append(col_width)
            if i < n - 1:
                widths.append(gutter)

        table = self.doc.add_table(rows=1 + depth * 2, cols=len(widths))
        clear_table_borders(table)
        fixed_layout(table, widths)

        head = table.rows[0]
        set_row_height(head, 20)
        for i, (heading, _) in enumerate(columns):
            cell = head.cells[i * 2]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(heading)
            set_run_font(run, BODY_FONT, Pt(10), bold=True)
            shade_cell(cell, HEADER_FILL)
            set_cell_borders(cell, top=(6, "8C8C8C"), left=(6, "8C8C8C"),
                             bottom=(6, "8C8C8C"), right=(6, "8C8C8C"))

        for f_idx in range(depth):
            line_row = table.rows[1 + f_idx * 2]
            label_row = table.rows[2 + f_idx * 2]
            set_row_height(line_row, 34)
            set_row_height(label_row, 12)
            for i, (_, fields) in enumerate(columns):
                col = i * 2
                line_cell = line_row.cells[col]
                label_cell = label_row.cells[col]
                line_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                label_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                if f_idx < len(fields):
                    set_cell_borders(line_cell, bottom=(6, "8C8C8C"))
                    run = label_cell.paragraphs[0].add_run(fields[f_idx])
                    set_run_font(run, BODY_FONT, Pt(8.5), color=FURNITURE_GREY)
        self.spacer(6)
        return table

    def note_box(self, heading, body, fill=NOTE_FILL):
        table = self.doc.add_table(rows=1, cols=1)
        clear_table_borders(table)
        fixed_layout(table, [CONTENT_WIDTH])
        cell = table.rows[0].cells[0]
        set_cell_borders(cell, top=(8, "8C8C8C"), left=(8, "8C8C8C"),
                         bottom=(8, "8C8C8C"), right=(8, "8C8C8C"))
        shade_cell(cell, fill)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(heading + " ")
        set_run_font(run, BODY_FONT, Pt(10.5), bold=True)
        self._emit_runs(para, body, size=Pt(10.5))
        self.spacer(8)
        return table

    # -- close out ---------------------------------------------------------

    def finish(self, filename, extra_disclaimer=None):
        self.spacer(10)
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        self._rule_under(para)
        run = para.add_run("")
        set_run_font(run, BODY_FONT, Pt(2))
        disclaimer = self.doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = DISCLAIMER
        if extra_disclaimer:
            text += " " + extra_disclaimer
        run = disclaimer.add_run(text)
        set_run_font(run, BODY_FONT, Pt(9.5), italic=True, color=FURNITURE_GREY)

        os.makedirs(OUT_DIR, exist_ok=True)
        path = os.path.join(OUT_DIR, filename)
        self.doc.save(path)
        return path


# ---------------------------------------------------------------------------
# 2.1  Subcontractor Agreement
# ---------------------------------------------------------------------------


def build_subcontractor_agreement():
    d = BinderDoc("Subcontractor Agreement")
    d.h1("SUBCONTRACTOR AGREEMENT")

    d.h2("PROJECT INFORMATION")
    d.fields([["Project Address"], ["Legal Description"]])

    d.h3("Owner / General Contractor")
    d.fields([["Name"], ["Address"], ["City / State / ZIP"], ["Phone", "Email"]])

    d.h3("Subcontractor")
    d.fields([
        ["Company Name"],
        ["License #", "License State"],
        ["Contact Person"],
        ["Address"],
        ["City / State / ZIP"],
        ["Phone", "Email"],
    ])

    d.h3("Project Description")
    d.ruled(2)

    d.h3("Key Dates")
    d.fields([
        ["Agreement Date", "Work Start Date"],
        ["Substantial Completion", "Final Completion"],
    ])

    d.page_break()

    d.h2("SCOPE OF WORK")
    d.label("Detailed description of work to be performed:")
    d.ruled(9)

    d.h3("Specifications and Standards")
    d.checklist([
        "Work shall comply with all applicable building codes and regulations.",
        "Work shall comply with approved plans dated [PLAN DATE].",
        "Work shall meet industry standards for [TRADE / STANDARD].",
        "Other specifications: [DESCRIBE].",
    ])

    d.h3("Materials")
    d.label("Materials to be provided by Subcontractor:")
    d.ruled(3)
    d.label("Materials to be provided by Owner:")
    d.ruled(3)

    d.h3("Exclusions (work NOT included in this agreement)")
    d.ruled(4)

    d.page_break()

    d.h2("PAYMENT TERMS")
    d.fields([["Total Contract Amount ($)"]])

    d.h3("Payment Schedule")
    d.grid(
        headers=["Draw #", "Milestone / Completion Point", "Amount ($)"],
        rows=[[str(i), "", ""] for i in range(1, 6)],
        widths=[0.8, 4.4, 1.8],
        align_right=(2,),
    )

    d.fields([
        ["Retention Held (%)", "Retention Amount ($)"],
        ["Payment Due Within (days of invoice)"],
    ])

    d.h3("Final Payment Conditions")
    d.p("Final payment shall be made upon:")
    d.checklist([
        "Completion of all work per specifications",
        "Passing final inspection",
        "Receipt of unconditional lien waiver from Subcontractor",
        "Receipt of lien waivers from all material suppliers",
        "Submission of warranties and operating manuals (if applicable)",
        "Correction of any punch list items",
    ])

    d.h3("Change Orders")
    d.p(
        "All changes to the scope of work must be approved in writing via Change "
        "Order Form before work begins. No additional payment will be made for "
        "work performed without a signed change order."
    )
    d.p("Change orders must include:")
    d.bullets([
        "Detailed description of change",
        "Cost impact (materials and labor)",
        "Schedule impact",
        "Signatures of both parties",
    ])

    d.page_break()

    d.h2("LEGAL REQUIREMENTS AND RESPONSIBILITIES")

    d.h3("Insurance Requirements")
    d.p(
        "Subcontractor shall maintain the following insurance coverage throughout "
        "the duration of this agreement:"
    )
    d.checklist([
        "General Liability Insurance: minimum $[AMOUNT] per occurrence",
        "Workers' Compensation Insurance (if required by state law)",
        "Vehicle Insurance: minimum $[AMOUNT] per occurrence",
    ])
    d.p(
        "Subcontractor shall provide a Certificate of Insurance to Owner before "
        "work begins."
    )

    d.h3("License Verification")
    d.checklist([
        "Subcontractor license verified with state licensing board",
        "License in good standing as of [DATE VERIFIED]",
    ])

    d.h3("Lien Waivers")
    d.p("Subcontractor agrees to provide:")
    d.bullets([
        "Conditional lien waiver with each progress payment request",
        "Unconditional lien waiver upon receipt of each progress payment",
        "Final unconditional lien waiver upon receipt of final payment",
        "Lien waivers from all material suppliers and sub-subcontractors",
    ])

    d.h3("Warranty")
    d.p("Subcontractor warrants that all work shall be:")
    d.bullets([
        "Performed in a workmanlike manner",
        "Free from defects in materials and workmanship",
        "In compliance with all applicable codes and regulations",
        "Performed by qualified personnel",
    ])
    d.fields([["Warranty Period (from substantial completion)"]])
    d.p(
        "Subcontractor shall promptly correct any defects in materials or "
        "workmanship that appear during the warranty period at no additional cost "
        "to Owner."
    )

    d.h3("Permits and Inspections")
    d.checklist([
        "Owner responsible for obtaining permits",
        "Subcontractor responsible for obtaining permits",
    ])
    d.p(
        "Subcontractor shall notify Owner at least 48 hours before inspections are "
        "required and shall coordinate with the inspector to schedule inspections. "
        "Subcontractor shall correct any work that fails inspection at no "
        "additional cost to Owner."
    )

    d.page_break()

    d.h2("DISPUTE RESOLUTION")
    d.p("Any disputes arising under this agreement shall be resolved as follows:")
    d.numbered([
        "Direct Communication: Parties shall first attempt to resolve disputes "
        "through direct communication within 7 days of the dispute arising.",
        "Written Notice: If the dispute is not resolved, the aggrieved party shall "
        "provide written notice detailing the dispute within 14 days.",
        "Meeting: Parties shall meet in person within 14 days of written notice to "
        "attempt resolution.",
        "Mediation: If the meeting does not resolve the dispute, parties agree to "
        "participate in mediation with a mutually agreed upon mediator. Cost of "
        "mediation shall be split equally.",
        "Arbitration / Litigation: If mediation fails, parties may pursue the "
        "option checked below.",
    ])
    d.checklist([
        "Binding arbitration",
        "Litigation in the courts of [COUNTY] County, [STATE]",
    ], indent=0.25)

    d.h3("Termination")
    d.p(
        "By Owner: Owner may terminate this agreement for cause (failure to "
        "perform, safety violations, code violations) with 7 days written notice. "
        "Owner may terminate without cause with 14 days written notice."
    )
    d.p(
        "By Subcontractor: Subcontractor may terminate for non-payment after "
        "providing 14 days written notice and opportunity to cure."
    )
    d.p(
        "Upon termination, Subcontractor shall be paid for all work completed to "
        "date, less any costs incurred by Owner to correct deficiencies."
    )

    d.h3("Indemnification")
    d.p(
        "Subcontractor agrees to indemnify and hold harmless Owner from any "
        "claims, damages, losses, or expenses (including attorney fees) arising "
        "from:"
    )
    d.bullets([
        "Subcontractor's negligent acts or omissions",
        "Injury to persons or property caused by Subcontractor or its employees",
        "Failure to comply with applicable laws and regulations",
        "Claims by Subcontractor's employees, suppliers, or sub-subcontractors",
    ])

    d.page_break()

    d.h2("ADDITIONAL TERMS AND CONDITIONS")

    d.h3("Site Access and Scheduling")
    d.p(
        "Subcontractor shall coordinate the work schedule with Owner and other "
        "trades."
    )
    d.fields([["Site Access Hours"]])

    d.h3("Cleanup")
    d.p(
        "Subcontractor shall maintain a clean and safe work area and remove all "
        "debris and materials at the end of each workday. Final cleanup shall "
        "include removal of all tools, equipment, and materials."
    )

    d.h3("Safety")
    d.p(
        "Subcontractor shall comply with all OSHA regulations and maintain a safe "
        "work environment. Subcontractor is responsible for the safety of its "
        "employees and shall provide all required personal protective equipment."
    )

    d.h3("Damage to Property")
    d.p(
        "Subcontractor shall be responsible for any damage to existing structures, "
        "fixtures, or property caused by Subcontractor or its employees."
    )

    d.h3("Assignment")
    d.p(
        "This agreement may not be assigned without written consent of both "
        "parties."
    )

    d.h3("Entire Agreement")
    d.p(
        "This agreement constitutes the entire agreement between the parties and "
        "supersedes all prior negotiations, representations, or agreements. This "
        "agreement may only be modified by written amendment signed by both "
        "parties."
    )

    d.h3("Governing Law")
    d.p(
        "This agreement shall be governed by the laws of the State of [STATE]."
    )

    d.h2("SIGNATURES")
    d.p(
        "By signing below, the parties acknowledge they have read, understood, and "
        "agree to all terms and conditions of this agreement."
    )
    d.signature_block([
        ("OWNER / GENERAL CONTRACTOR",
         ["Signature", "Printed Name", "Title", "Date"]),
        ("SUBCONTRACTOR",
         ["Signature", "Printed Name", "Title", "Date"]),
    ])

    d.h3("Witness (optional)")
    d.signature_block([
        ("WITNESS", ["Signature", "Printed Name", "Date"]),
    ])

    return d.finish("2.1-subcontractor-agreement-template.docx")


# ---------------------------------------------------------------------------
# 2.2  Change Order Form
# ---------------------------------------------------------------------------


def build_change_order_form():
    d = BinderDoc("Change Order Form")
    d.h1("CHANGE ORDER FORM")

    d.fields([["Change Order Number", "Date"]])

    d.h2("PROJECT INFORMATION")
    d.fields([
        ["Project Address"],
        ["Owner Name"],
        ["Subcontractor / Trade"],
        ["Original Contract Date"],
    ])

    d.h2("FINANCIAL SUMMARY")
    d.grid(
        headers=["Description", "Amount ($)"],
        rows=[
            ["Original contract amount", ""],
            ["Previous change orders (total)", ""],
            ["Revised contract amount (before this change order)", ""],
            ["THIS change order amount", ""],
            ["New contract amount (including this change order)", ""],
        ],
        widths=[4.9, 2.1],
        bold_rows={4},
        align_right=(1,),
    )
    d.p("This change order is an:")
    d.checklist(["Addition (+)", "Deduction (−)"])

    d.h2("DESCRIPTION OF CHANGE")
    d.label("Detailed description of work to be added, deleted, or modified:")
    d.ruled(10)

    d.page_break()

    d.h2("REASON FOR CHANGE")
    d.p("Check all that apply:")
    d.checklist([
        "Owner-requested modification",
        "Design change",
        "Unforeseen site conditions",
        "Code requirement / inspector request",
        "Material substitution or unavailability",
        "Correction of error or omission",
        "Upgrade / enhancement",
        "Value engineering",
        "Other: [DESCRIBE]",
    ])
    d.label("Detailed explanation:")
    d.ruled(5)

    d.h2("COST BREAKDOWN")
    d.grid(
        headers=["Description", "Quantity", "Cost ($)"],
        rows=[
            ["LABOR", "", ""],
            ["", "", ""],
            ["", "", ""],
            ["", "", ""],
            ["Labor subtotal", "", ""],
            ["MATERIALS", "", ""],
            ["", "", ""],
            ["", "", ""],
            ["", "", ""],
            ["", "", ""],
            ["Materials subtotal", "", ""],
            ["EQUIPMENT / OTHER", "", ""],
            ["", "", ""],
            ["", "", ""],
            ["Equipment subtotal", "", ""],
            ["TOTAL COST IMPACT", "", ""],
        ],
        widths=[4.0, 1.4, 1.6],
        bold_rows={0, 4, 5, 10, 11, 14, 15},
        align_right=(2,),
    )

    d.page_break()

    d.h2("SCHEDULE IMPACT")
    d.fields([
        ["Original Completion Date", "New Completion Date"],
        ["Schedule Impact (+/− days)"],
    ])
    d.label("Explanation of schedule impact:")
    d.ruled(3)

    d.h2("IMPACT ON OTHER TRADES")
    d.checklist([
        "No impact on other trades",
        "Impacts other trades (explain below)",
    ])
    d.ruled(3)
    d.fields([["Trades Affected"]])

    d.h2("SUPPORTING DOCUMENTATION")
    d.p("Attached documents (check all that apply):")
    d.checklist([
        "Revised drawings / sketches",
        "Material quotes / invoices",
        "Photos of conditions requiring change",
        "Inspector notes / requirements",
        "Engineer specifications",
        "Other: [DESCRIBE]",
    ])

    d.h2("PAYMENT TERMS FOR THIS CHANGE ORDER")
    d.checklist([
        "Payment due with next regular draw",
        "Payment upon completion of change order work",
        "Payment in installments: [DESCRIBE SCHEDULE]",
        "Other: [DESCRIBE]",
    ])

    d.h2("APPROVALS AND SIGNATURES")
    d.p(
        "By signing below, both parties agree to the changes described in this "
        "Change Order and acknowledge that this Change Order modifies the original "
        "contract accordingly."
    )
    d.signature_block([
        ("OWNER APPROVAL",
         ["Signature", "Printed Name", "Company", "Date"]),
        ("SUBCONTRACTOR APPROVAL",
         ["Signature", "Printed Name", "Company", "Date"]),
    ])

    d.h2("NOTES")
    d.ruled(4)

    return d.finish("2.2-change-order-form.docx")


# ---------------------------------------------------------------------------
# 2.3  Lien Waiver Templates (four variants)
# ---------------------------------------------------------------------------

CLAIMANT_FIELDS = [
    ["Name of Claimant"],
    ["Company Name"],
    ["License Number"],
    ["Address"],
    ["Phone"],
]

# One notice per variant. The source PDF prints the unconditional warning on all
# four forms, which is self-contradictory on the two conditional ones.
CONDITIONAL_NOTICE = (
    "This waiver is CONDITIONAL. It becomes effective only when the payment "
    "identified below is actually received by the person signing. If that payment "
    "is not received, this document waives nothing. To acknowledge a payment "
    "already received, use an unconditional waiver instead."
)

UNCONDITIONAL_WARNING = (
    "This document waives rights unconditionally and states that you have been "
    "paid for giving up those rights. This document is enforceable against you if "
    "you sign it, even if you have not been paid. If you have not been paid, use a "
    "conditional waiver and release form."
)

FINAL_CONDITIONAL_NOTICE = (
    "This waiver is CONDITIONAL. It waives all remaining lien rights only when the "
    "FINAL payment identified below is actually received by the person signing. If "
    "that payment is not received, this document waives nothing. To acknowledge "
    "final payment already received, use an unconditional final waiver instead."
)

FINAL_UNCONDITIONAL_WARNING = (
    "This document waives ALL rights unconditionally and states that you have been "
    "paid IN FULL for giving up those rights. This document is enforceable against "
    "you if you sign it, even if you have not been paid in full. If you have not "
    "been paid in full, use a conditional waiver and release form."
)

LIEN_WAIVER_DISCLAIMER = (
    "Several states prescribe the exact wording of lien waivers; check your "
    "state's statute before signing."
)


def _waiver_header(d, title, subtitle):
    d.h1(title)
    d.subtitle(subtitle)
    d.fields([
        ["Project Name"],
        ["Project Address"],
        ["Owner Name"],
        ["Date"],
    ])


def _claimant_and_signature(d):
    d.h2("CLAIMANT INFORMATION")
    d.fields(CLAIMANT_FIELDS)
    d.signature_block([
        ("CLAIMANT", ["Signature", "Printed Name", "Title", "Date"]),
    ])


def build_lien_waivers():
    d = BinderDoc("Lien Waiver Templates")

    # ---- 1. Conditional waiver, progress payment -------------------------
    _waiver_header(d, "CONDITIONAL LIEN WAIVER", "Progress / Partial Payment")

    d.p(
        "Upon receipt of payment in the amount of $[PAYMENT AMOUNT], the "
        "undersigned hereby waives and releases any and all lien rights, stop "
        "payment notice rights, and bond rights the undersigned has or may have "
        "against the above-referenced property for labor, services, equipment, or "
        "materials furnished to the property through [THROUGH-DATE]."
    )
    d.p(
        "This waiver is CONDITIONAL upon payment. This waiver and release is void "
        "and of no effect unless and until the undersigned actually receives "
        "payment in the amount stated above.",
        bold=True,
    )

    d.h2("PAYMENT INFORMATION")
    d.fields([
        ["Payment Amount ($)", "Check Number"],
        ["Date of Check"],
        ["Payment Period From", "Payment Period To"],
    ])

    d.h2("EXCEPTIONS")
    d.p("This waiver does not cover:")
    d.bullets([
        "Any work, materials, or services provided after the through-date listed "
        "above",
        "Any retention amounts held",
        "Disputed claims for extra work in the amount of $[AMOUNT]",
    ])
    d.label("Other exceptions:")
    d.ruled(2)

    _claimant_and_signature(d)
    d.note_box("NOTICE:", CONDITIONAL_NOTICE, fill=NOTICE_FILL)

    d.page_break()

    # ---- 2. Unconditional waiver, progress payment -----------------------
    _waiver_header(d, "UNCONDITIONAL LIEN WAIVER", "Progress / Partial Payment")

    d.p(
        "The undersigned has been paid and has received progress payment in the "
        "amount of $[PAYMENT AMOUNT], and hereby waives and releases any and all "
        "lien rights, stop payment notice rights, and bond rights the undersigned "
        "has or may have against the above-referenced property for labor, "
        "services, equipment, or materials furnished to the property through "
        "[THROUGH-DATE]."
    )
    d.p(
        "This waiver is UNCONDITIONAL. This waiver and release is effective "
        "immediately upon signing and is not dependent upon receipt of payment.",
        bold=True,
    )

    d.h2("PAYMENT INFORMATION")
    d.fields([
        ["Payment Amount Received ($)", "Check Number"],
        ["Date Payment Received"],
        ["Payment Period From", "Payment Period To"],
    ])

    d.h2("AMOUNT SUMMARY")
    d.grid(
        headers=["Description", "Amount ($)"],
        rows=[
            ["Total contract amount to date", ""],
            ["Previous payments received", ""],
            ["This payment amount", ""],
            ["Balance remaining", ""],
        ],
        widths=[4.9, 2.1],
        bold_rows={3},
        align_right=(1,),
    )

    d.h2("EXCEPTIONS")
    d.p("This waiver does not cover:")
    d.bullets([
        "Any work, materials, or services provided after the through-date listed "
        "above",
        "Any retention amounts held, in the amount of $[AMOUNT]",
        "Disputed claims for extra work in the amount of $[AMOUNT]",
    ])
    d.label("Other exceptions:")
    d.ruled(2)

    _claimant_and_signature(d)
    d.note_box("WARNING:", UNCONDITIONAL_WARNING)

    d.page_break()

    # ---- 3. Conditional waiver, final payment ----------------------------
    _waiver_header(d, "CONDITIONAL LIEN WAIVER", "Final Payment")

    d.p(
        "Upon receipt of FINAL payment in the amount of $[FINAL PAYMENT AMOUNT], "
        "the undersigned hereby waives and releases any and all lien rights, stop "
        "payment notice rights, and bond rights the undersigned has or may have "
        "against the above-referenced property for all labor, services, equipment, "
        "or materials furnished to the property."
    )
    d.p(
        "This waiver is CONDITIONAL upon receipt of final payment. This waiver and "
        "release is void and of no effect unless and until the undersigned "
        "actually receives final payment in the amount stated above.",
        bold=True,
    )

    d.h2("FINAL PAYMENT INFORMATION")
    d.grid(
        headers=["Description", "Amount ($)"],
        rows=[
            ["Total contract amount", ""],
            ["Previous payments received", ""],
            ["Final payment amount", ""],
            ["Retention amount being released", ""],
        ],
        widths=[4.9, 2.1],
        bold_rows={2},
        align_right=(1,),
    )
    d.fields([["Check Number", "Date of Check"]])

    d.h2("CERTIFICATION")
    d.p("The undersigned certifies that:")
    d.checklist([
        "All work has been completed per contract specifications",
        "All materials have been paid for",
        "All subcontractors and suppliers have been paid in full",
        "All required warranties have been provided to Owner",
        "All punch list items have been completed",
    ])

    d.h2("EXCEPTIONS")
    d.p("This waiver does not cover (if none, write “NONE”):")
    d.ruled(3)

    _claimant_and_signature(d)
    d.note_box("NOTICE:", FINAL_CONDITIONAL_NOTICE, fill=NOTICE_FILL)

    d.page_break()

    # ---- 4. Unconditional waiver, final payment --------------------------
    _waiver_header(d, "UNCONDITIONAL LIEN WAIVER", "Final Payment")

    d.p(
        "The undersigned has been paid in full and has received FINAL payment in "
        "the amount of $[FINAL PAYMENT AMOUNT], and hereby waives and releases any "
        "and all lien rights, stop payment notice rights, and bond rights the "
        "undersigned has or may have against the above-referenced property for all "
        "labor, services, equipment, or materials furnished to the property."
    )
    d.p(
        "This waiver is UNCONDITIONAL and FINAL. This waiver and release is "
        "effective immediately upon signing and releases all rights to file liens "
        "or claims against the property.",
        bold=True,
    )

    d.h2("FINAL PAYMENT INFORMATION")
    d.grid(
        headers=["Description", "Amount ($)"],
        rows=[
            ["Total contract amount", ""],
            ["Total change orders", ""],
            ["Final contract amount", ""],
            ["Previous payments received", ""],
            ["Retention released", ""],
            ["Final payment amount", ""],
        ],
        widths=[4.9, 2.1],
        bold_rows={2, 5},
        align_right=(1,),
    )
    d.fields([["Check Number", "Date Payment Received"]])

    d.h2("FINAL CERTIFICATION")
    d.p("The undersigned certifies that:")
    d.checklist([
        "All work has been completed per contract specifications",
        "All materials have been paid for in full",
        "All subcontractors and suppliers have been paid in full",
        "All required warranties have been provided to Owner",
        "All punch list items have been completed",
        "Final inspection has been passed",
        "No outstanding claims or disputes exist",
    ])

    d.h2("EXCEPTIONS")
    d.p(
        "This is a FINAL waiver with NO exceptions. If any claims remain, do NOT "
        "sign this form.",
        bold=True,
    )

    _claimant_and_signature(d)
    d.note_box("WARNING:", FINAL_UNCONDITIONAL_WARNING)

    return d.finish("2.3-lien-waiver-templates.docx",
                    extra_disclaimer=LIEN_WAIVER_DISCLAIMER)


# ---------------------------------------------------------------------------
# 2.4  Payment Draw Schedule
# ---------------------------------------------------------------------------


def _draw_block(d, number, final=False):
    heading = "DRAW %d — FINAL PAYMENT" % number if final else "DRAW %d" % number
    d.h2(heading)

    d.fields([["Draw Date", "Percentage Complete (%)"]])
    if final:
        d.p("The final draw closes the contract at 100% complete.", italic=True,
            color=FURNITURE_GREY)

    d.label("Work completed this period:")
    d.ruled(3)

    retention_label = ("Plus retention release" if final
                       else "Less retention (per contract %)")
    balance_value = "$0.00" if final else ""

    d.grid(
        headers=["Description", "Amount ($)", "Notes"],
        rows=[
            ["Labor", "", ""],
            ["Materials", "", ""],
            ["Equipment", "", ""],
            ["Subtotal this draw", "", ""],
            [retention_label, "", ""],
            ["Amount due this draw", "", ""],
            ["Previous payments", "", ""],
            ["Total paid to date", "", ""],
            ["Contract balance remaining", balance_value, ""],
        ],
        widths=[2.8, 1.7, 2.5],
        bold_rows={3, 5, 7, 8},
        align_right=(1,),
    )

    d.h3("Lien Waivers Received")
    d.checklist([
        "Conditional lien waiver submitted",
        "Unconditional lien waiver for previous payment",
        "Supplier lien waivers attached",
    ])

    if final:
        d.h3("Final Payment Requirements — ALL must be checked")
        d.checklist([
            "Unconditional final lien waiver received",
            "All supplier and subcontractor lien waivers received",
            "Final inspection passed",
            "All punch list items completed",
            "Warranties and manuals received",
            "All required permits closed out",
        ])

    d.h3("Inspection Status")
    d.checklist([
        "Inspection passed",
        "Inspection not yet required",
        "Inspection failed (see notes)",
    ])
    d.fields([
        ["Inspector", "Inspection Date"],
        ["Payment Date", "Check #"],
    ])

    d.signature_block([
        ("OWNER-BUILDER APPROVAL", ["Signature", "Printed Name", "Date"]),
    ])


def build_payment_draw_schedule():
    d = BinderDoc("Payment Draw Schedule")
    d.h1("PAYMENT DRAW SCHEDULE")

    d.h2("PROJECT AND CONTRACT INFORMATION")
    d.fields([
        ["Project Name"],
        ["Project Address"],
        ["Owner-Builder"],
        ["Subcontractor / Trade"],
        ["Original Contract Amount ($)", "Retention (%)"],
    ])

    d.p(
        "Record every draw below. Do not release a payment until the matching lien "
        "waiver is in hand and any required inspection has passed.",
        italic=True,
        color=FURNITURE_GREY,
    )

    for number in range(1, 5):
        d.page_break()
        _draw_block(d, number)

    d.page_break()
    _draw_block(d, 5, final=True)

    return d.finish("2.4-payment-draw-schedule.docx")


# ---------------------------------------------------------------------------
# Verification
# ---------------------------------------------------------------------------


def iter_docx_text(path):
    """Yield (location, text) for every paragraph in body, tables and footers."""
    doc = Document(path)

    def walk_table(table, prefix):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    yield ("%s r%dc%d" % (prefix, r, c), para.text)
                for inner in cell.tables:
                    yield from walk_table(inner, "%s r%dc%d>" % (prefix, r, c))

    for i, para in enumerate(doc.paragraphs):
        yield ("body p%d" % i, para.text)
    for t, table in enumerate(doc.tables):
        yield from walk_table(table, "table%d" % t)
    for s, section in enumerate(doc.sections):
        for part_name in ("footer", "header"):
            part = getattr(section, part_name)
            for i, para in enumerate(part.paragraphs):
                yield ("section%d %s p%d" % (s, part_name, i), para.text)


def iter_block_items(doc):
    """Yield ('p'|'tbl', text) for every top-level block, in document order."""
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield "tbl", Table(child, doc)


# (title, subtitle, notice heading, notice body) in the order they must appear.
WAIVER_VARIANTS = [
    ("CONDITIONAL LIEN WAIVER", "Progress / Partial Payment",
     "NOTICE:", CONDITIONAL_NOTICE),
    ("UNCONDITIONAL LIEN WAIVER", "Progress / Partial Payment",
     "WARNING:", UNCONDITIONAL_WARNING),
    ("CONDITIONAL LIEN WAIVER", "Final Payment",
     "NOTICE:", FINAL_CONDITIONAL_NOTICE),
    ("UNCONDITIONAL LIEN WAIVER", "Final Payment",
     "WARNING:", FINAL_UNCONDITIONAL_WARNING),
]


def verify_lien_waiver_warnings(path):
    """Each waiver variant must carry its own notice and no other variant's.

    The source PDF prints the unconditional warning on all four forms, which
    tells a signer of a conditional waiver that they have already been paid.
    This check exists so that defect cannot creep back in.
    """
    name = os.path.basename(path)
    problems = []
    doc = Document(path)

    segments = []
    current = None
    for kind, block in iter_block_items(doc):
        if kind == "p":
            if block.style.name == "Heading 1":
                current = {"title": block.text.strip(), "subtitle": None,
                           "text": []}
                segments.append(current)
                continue
            if current is not None:
                if current["subtitle"] is None and block.text.strip():
                    current["subtitle"] = block.text.strip()
                current["text"].append(block.text)
        elif current is not None:
            current["text"].append(
                "\n".join(cell.text for row in block.rows for cell in row.cells))

    if len(segments) != len(WAIVER_VARIANTS):
        problems.append("%s: found %d waiver variants, expected %d"
                        % (name, len(segments), len(WAIVER_VARIANTS)))
        return problems

    all_notices = [(heading, body) for _, _, heading, body in WAIVER_VARIANTS]
    for idx, (title, subtitle, heading, body) in enumerate(WAIVER_VARIANTS):
        segment = segments[idx]
        joined = "\n".join(segment["text"])
        where = "%s: variant %d (%s / %s)" % (name, idx + 1, title, subtitle)
        if segment["title"] != title:
            problems.append("%s: title is %r" % (where, segment["title"]))
        if segment["subtitle"] != subtitle:
            problems.append("%s: subtitle is %r" % (where, segment["subtitle"]))
        if body not in joined:
            problems.append("%s: missing its %s notice" % (where, heading))
        for other_heading, other_body in all_notices:
            if other_body is not body and other_body in joined:
                problems.append("%s: carries the wrong notice (%s)"
                                % (where, other_heading))

    if LIEN_WAIVER_DISCLAIMER not in "\n".join(
            p.text for p in doc.paragraphs):
        problems.append("%s: state-statute disclaimer missing" % name)
    return problems


def verify_docx(path):
    """Re-open a generated document and assert the house rules hold."""
    name = os.path.basename(path)
    problems = []
    texts = list(iter_docx_text(path))
    joined = "\n".join(t for _, t in texts)

    for where, text in texts:
        if "____" in text:
            problems.append("%s: underscore run in %s -> %r" % (name, where, text[:60]))
        if "© 2024" in text or "(c) 2024" in text:
            problems.append("%s: stale copyright in %s" % (name, where))
        if "/Users/" in text or "file:///" in text or ".html" in text:
            problems.append("%s: filesystem path leaked in %s" % (name, where))
        if "Owner-Builder Job Site Binder | Section" in text:
            problems.append("%s: PDF running header leaked in %s" % (name, where))

    if COPYRIGHT not in joined:
        problems.append("%s: footer copyright missing" % name)
    if DISCLAIMER not in joined:
        problems.append("%s: attorney-review disclaimer missing" % name)

    doc = Document(path)
    if doc.styles["Normal"].font.name != BODY_FONT:
        problems.append("%s: body font is not %s" % (name, BODY_FONT))
    if doc.styles["Normal"].font.size != BODY_SIZE:
        problems.append("%s: body size is not 11pt" % name)

    italic_disclaimer = False
    for para in doc.paragraphs:
        if DISCLAIMER in para.text:
            italic_disclaimer = any(run.font.italic for run in para.runs)
    if not italic_disclaimer:
        problems.append("%s: disclaimer is not italic" % name)

    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    if not headings:
        problems.append("%s: no Heading 1 title" % name)

    stats = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "h1": len(headings),
        "bordered_cells": sum(
            1
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
            if cell._tc.tcPr is not None
            and cell._tc.tcPr.find(qn("w:tcBorders")) is not None
        ),
    }
    return problems, stats


# ---------------------------------------------------------------------------


BUILDERS = [
    build_subcontractor_agreement,
    build_change_order_form,
    build_lien_waivers,
    build_payment_draw_schedule,
]


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    paths = [build() for build in BUILDERS]

    print("WORD DOCUMENTS")
    print("=" * 78)
    all_problems = []
    for path in paths:
        problems, stats = verify_docx(path)
        if os.path.basename(path).startswith("2.3-"):
            problems = problems + verify_lien_waiver_warnings(path)
        all_problems += problems
        size = os.path.getsize(path)
        status = "PASS" if not problems else "FAIL"
        print("  [%s] %-46s %7s bytes" % (status, os.path.basename(path),
                                          "{:,}".format(size)))
        print("         paragraphs=%(paragraphs)d  tables=%(tables)d  "
              "bordered cells=%(bordered_cells)d" % stats)
    if all_problems:
        print("\n  PROBLEMS")
        for problem in all_problems:
            print("   - " + problem)
    else:
        print("\n  All checks passed: no underscore runs, no '© 2024', no leaked")
        print("  paths or PDF running headers; Calibri 11 body, Heading 1 titles,")
        print("  italic disclaimer and 2026 footer present in all 4 documents.")
    print("  Output: %s" % OUT_DIR)
    return 1 if all_problems else 0


if __name__ == "__main__":
    raise SystemExit(main())
